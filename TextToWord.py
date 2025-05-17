#!/usr/bin/env python
# coding: utf-8

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import argparse
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_word_document(text_file: str, output_file: str = None) -> None:
    """
    Convert a text file to a formatted Word document
    
    Args:
        text_file (str): Path to input text file
        output_file (str, optional): Path to save the Word document. If not provided,
                                   will use the input filename with .docx extension
    """
    try:
        # Validate input file
        if not os.path.exists(text_file):
            raise FileNotFoundError(f"The file {text_file} does not exist")
            
        # Set default output filename if not provided
        if output_file is None:
            output_file = os.path.splitext(text_file)[0] + '.docx'
            
        # Create a new Document
        doc = Document()
        
        # Set default document properties
        doc.core_properties.title = os.path.basename(text_file)
        
        # Set default styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Read and process the text file
        with open(text_file, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Split content into paragraphs (handling different line endings)
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        # Add paragraphs to document with formatting
        for para_text in paragraphs:
            # Create new paragraph
            paragraph = doc.add_paragraph()
            
            # Add text with proper formatting
            paragraph.add_run(para_text)
            
            # Set paragraph formatting
            paragraph.paragraph_format.space_after = Pt(10)
            paragraph.paragraph_format.line_spacing = 1.15
            
            # If paragraph is short (potential heading), make it bold
            if len(para_text.split()) <= 10 and not para_text.endswith('.'):
                paragraph.runs[0].bold = True
                paragraph.paragraph_format.space_before = Pt(12)
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Save the document
        doc.save(output_file)
        logger.info(f"Successfully created Word document: {output_file}")
        
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        raise

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Convert text file to Word document')
    parser.add_argument('input_file', help='Path to input text file')
    parser.add_argument('--output', '-o', help='Path to output Word document (optional)')
    
    args = parser.parse_args()
    
    try:
        create_word_document(args.input_file, args.output)
    except Exception as e:
        logger.error(f"Error: {str(e)}")
        exit(1) 